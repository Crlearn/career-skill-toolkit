#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简历优化工具：读取简历文件 + 生成 HTML 简历 + 版本管理 + 字段映射

用法:
  python generate_resume.py read <file_path>          读取 PDF/Word/TXT 简历，输出纯文本 JSON
  python generate_resume.py build <json_file> <output> 读取结构化 JSON，生成 HTML 简历文件
      可选: --pages 1|2  指定简历页数（默认1页，2页时内容更宽松）
  python generate_resume.py init [output_path]         生成空白简历 JSON 模板
  python generate_resume.py setup                     一键初始化（检查依赖+创建模板文件）
  python generate_resume.py map <json> [output]       字段映射：将非标准字段名映射为标准模板字段
  python generate_resume.py version save <json> <name> 保存当前简历版本
  python generate_resume.py version list               列出所有已保存版本
  python generate_resume.py version diff <v1> <v2>     对比两个版本的差异
  python generate_resume.py version restore <name>     回滚到指定版本

字段映射说明:
  build 命令会自动执行字段映射，将"姓名"→name、"手机"→phone 等非标准字段名
  转换为标准模板字段。也可用 map 命令单独执行映射，查看映射结果后再 build。
"""

import json
import os
import sys
import base64
import html as html_mod


def configure_output_encoding():
    for stream in (sys.stdout, sys.stderr):
        reconfigure = getattr(stream, "reconfigure", None)
        if callable(reconfigure):
            try:
                reconfigure(encoding="utf-8", errors="replace")
            except Exception:
                pass


configure_output_encoding()


# ---------------------------------------------------------------------------
# 空白简历模板
# ---------------------------------------------------------------------------

BLANK_TEMPLATE = {
    "name": "",
    "gender": "",
    "birth_date": "",
    "current_city": "",
    "domicile_city": "",
    "phone": "",
    "email": "",
    "objective": "",
    "all_objectives": [],
    "target_industry": "",
    "target_city": "",
    "expected_salary": "",
    "job_type": "",
    "photo": "",
    "core_advantages": [],
    "education": [
        {
            "school": "",
            "degree": "",
            "major": "",
            "period": "",
            "school_tag": "",
            "gpa": "",
            "courses": []
        }
    ],
    "experience": [
        {
            "section": "项目经历",
            "company": "",
            "role": "",
            "period": "",
            "tags": [],
            "bullets": [
                {"label": "", "text": ""}
            ]
        }
    ],
    "skills": [
        {"label": "", "text": ""}
    ],
    "awards": [
        {"name": "", "level": "", "time": ""}
    ],
    "portfolio": [
        {"title": "", "url": ""}
    ],
    "self_evaluation": "",
    "all_self_evaluations": {}
}


# ---------------------------------------------------------------------------
# JSON 校验
# ---------------------------------------------------------------------------

REQUIRED_FIELDS = ["name", "phone", "email", "objective"]


def validate_resume(data: dict, strict: bool = True) -> list:
    """校验简历 JSON 必填字段，返回错误列表。
    strict=True: 缺必填字段报错（build 时使用）
    strict=False: 缺必填字段仅警告（map 模式使用）
    """
    errors = []
    warnings = []
    for field in REQUIRED_FIELDS:
        if not data.get(field):
            msg = f"缺少必填字段: {field}"
            if strict:
                errors.append(msg)
            else:
                warnings.append(msg)
    edu = data.get("education", [])
    if isinstance(edu, dict):
        edu = [edu]
    if not edu:
        if strict:
            errors.append("缺少教育经历")
        else:
            warnings.append("缺少教育经历")
    else:
        for i, e in enumerate(edu):
            for ef in ["school", "degree", "major", "period"]:
                if not e.get(ef):
                    msg = f"教育经历[{i}] 缺少必填字段: {ef}"
                    if strict:
                        errors.append(msg)
                    else:
                        warnings.append(msg)
    if strict:
        return errors
    return {"errors": errors, "warnings": warnings}


# ---------------------------------------------------------------------------
# 简历读取
# ---------------------------------------------------------------------------

def _clean_text(text: str) -> str:
    import re
    cleaned = re.sub(r'[\ue000-\uf8ff\u0000-\u001f]', '', text)
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    return cleaned.strip()


def read_pdf(file_path: str) -> dict:
    try:
        import fitz
        doc = fitz.open(file_path)
        pages = []
        empty_pages = []
        for i, page in enumerate(doc):
            page_text = page.get_text().strip()
            pages.append(page_text)
            if not page_text:
                empty_pages.append(i + 1)
        doc.close()

        text = "\n".join(pages)
        text = _clean_text(text)

        result = {
            "page_count": len(pages),
            "empty_pages": empty_pages if empty_pages else None,
        }

        if not text.strip():
            result["warning"] = "PDF 可能是扫描件（图片格式），无法提取文本。请手动复制简历内容粘贴。"
            return result

        if empty_pages:
            result["warning"] = (
                f"第 {', '.join(str(p) for p in empty_pages)} 页内容为空，"
                f"可能是图片或特殊格式。如果简历有内容遗漏，请手动补充。"
            )

        result["text"] = text
        return result
    except ImportError:
        return {"error": "缺少 PyMuPDF 库，请运行: pip install PyMuPDF"}
    except Exception as e:
        return {"error": f"PDF 读取失败: {str(e)}"}


def read_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        return "\n".join(paragraphs).strip()
    except ImportError:
        return ""
    except Exception:
        return ""


def read_txt(file_path: str) -> str:
    for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read().strip()
        except (UnicodeDecodeError, UnicodeError):
            continue
    return ""


def read_resume(file_path: str) -> dict:
    if not os.path.isfile(file_path):
        return {"error": f"文件不存在: {file_path}"}
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        pdf_result = read_pdf(file_path)
        if "error" in pdf_result:
            return pdf_result
        text = pdf_result.get("text", "")
        warning = pdf_result.get("warning")
        if not text:
            return {"error": pdf_result.get("warning") or "文件内容为空或无法解析，请尝试复制简历文本直接粘贴"}
        result = {
            "status": "success",
            "file_path": file_path,
            "file_type": ext,
            "text": text,
            "char_count": len(text),
        }
        if warning:
            result["warning"] = warning
        if pdf_result.get("page_count"):
            result["page_count"] = pdf_result["page_count"]
        return result
    elif ext == ".docx":
        text = read_docx(file_path)
    elif ext in (".txt", ".text", ".md"):
        text = read_txt(file_path)
    else:
        return {"error": f"不支持的文件格式: {ext}，支持 PDF / Word(.docx) / TXT"}
    if not text:
        return {"error": "文件内容为空或无法解析，请尝试复制简历文本直接粘贴"}
    text = _clean_text(text)
    return {
        "status": "success",
        "file_path": file_path,
        "file_type": ext,
        "text": text,
        "char_count": len(text),
    }


# ---------------------------------------------------------------------------
# HTML 简历生成 — 国内商务简历样式
# ---------------------------------------------------------------------------

HTML_TEMPLATE = r"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{name} - 简历</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
  font-family: "PingFang SC", "Microsoft YaHei", "Source Han Sans SC", sans-serif;
  color: #333;
  line-height: 1.55;
  background: #f0f0f0;
  font-size: 13px;
}}
.resume {{
  max-width: 210mm;
  margin: 20px auto;
  padding: 18mm 20mm 14mm;
  background: white;
  box-shadow: 0 1px 6px rgba(0,0,0,0.08);
}}
/* ---- Header ---- */
.header {{
  display: flex;
  justify-content: space-between;
  align-items: flex-start;
  padding-bottom: 8px;
  border-bottom: 1.5px solid #2c3e50;
  margin-bottom: 12px;
}}
.header-left h1 {{
  font-size: 22px;
  font-weight: 700;
  color: #1a1a1a;
  letter-spacing: 1px;
  margin-bottom: 4px;
}}
.header-contact {{
  font-size: 12px;
  color: #555;
  line-height: 1.7;
}}
.header-contact .label {{
  color: #888;
}}
.header-photo {{
  width: 108px;
  height: 144px;
  object-fit: cover;
  border-radius: 4px;
  border: 1px solid #ddd;
  flex-shrink: 0;
}}
/* ---- Core Advantages ---- */
.core-advantages {{
  margin-bottom: 10px;
  padding: 8px 12px;
  background: #f8f9fa;
  border-left: 3px solid #2c3e50;
  border-radius: 2px;
}}
.core-advantages-title {{
  font-size: 13px;
  font-weight: 700;
  color: #1a1a1a;
  margin-bottom: 4px;
}}
.core-advantages-list {{
  margin: 0;
  padding-left: 0;
  list-style: none;
}}
.core-advantages-list li {{
  position: relative;
  padding-left: 12px;
  margin-bottom: 2px;
  font-size: 12.5px;
  line-height: 1.55;
  color: #333;
}}
.core-advantages-list li::before {{
  content: "▸";
  position: absolute;
  left: 1px;
  color: #2c3e50;
  font-weight: bold;
}}
/* ---- Section ---- */
.section {{
  margin-bottom: 10px;
}}
.section-title {{
  font-size: 14px;
  font-weight: 700;
  color: #1a1a1a;
  border-bottom: 1px solid #ddd;
  padding-bottom: 3px;
  margin-bottom: 6px;
}}
/* ---- Education ---- */
.edu-row {{
  display: flex;
  justify-content: space-between;
  align-items: baseline;
  margin-bottom: 1px;
  font-size: 13px;
}}
.edu-left {{ display: flex; align-items: baseline; gap: 6px; }}
.edu-school {{ font-weight: 700; }}
.edu-tag {{
  display: inline-block;
  background: #e8f4fd;
  color: #2980b9;
  font-size: 10.5px;
  padding: 0 5px;
  border-radius: 2px;
  font-weight: 600;
}}
.edu-detail {{ color: #555; }}
.edu-date {{ color: #888; font-size: 12px; white-space: nowrap; }}
.edu-courses {{
  font-size: 12px;
  color: #666;
  margin-top: 1px;
  line-height: 1.6;
}}
/* ---- Experience ---- */
.exp-header {{
  display: flex;
  justify-content: space-between;
  align-items: baseline;
  margin-bottom: 4px;
}}
.exp-left {{ display: flex; align-items: baseline; gap: 8px; }}
.exp-company {{ font-weight: 700; font-size: 13px; }}
.exp-role {{ color: #555; font-size: 13px; }}
.exp-date {{ color: #888; font-size: 12px; white-space: nowrap; }}
.bullet-list {{
  margin: 0;
  padding-left: 0;
  list-style: none;
}}
.bullet-list li {{
  position: relative;
  padding-left: 12px;
  margin-bottom: 1px;
  font-size: 12.5px;
  line-height: 1.55;
  color: #444;
}}
.bullet-list li::before {{
  content: "\00b7";
  position: absolute;
  left: 1px;
  color: #888;
  font-weight: bold;
}}
.bullet-label {{
  font-weight: 600;
  color: #222;
}}
/* ---- Skills ---- */
.skill-item {{
  margin-bottom: 1px;
  font-size: 12.5px;
  line-height: 1.55;
  padding-left: 12px;
  position: relative;
  color: #444;
}}
.skill-item::before {{
  content: "\00b7";
  position: absolute;
  left: 1px;
  color: #888;
  font-weight: bold;
}}
.skill-label {{
  font-weight: 600;
  color: #222;
}}
/* ---- Awards ---- */
.award-row {{
  display: flex;
  justify-content: space-between;
  align-items: baseline;
  margin-bottom: 1px;
  font-size: 12.5px;
  color: #444;
}}
.award-name {{ font-weight: 600; color: #222; }}
.award-level {{ color: #555; margin-left: 4px; }}
.award-time {{ color: #888; font-size: 12px; white-space: nowrap; }}
/* ---- Portfolio ---- */
.portfolio-item {{
  margin-bottom: 1px;
  font-size: 12.5px;
  line-height: 1.55;
  padding-left: 12px;
  position: relative;
  color: #444;
}}
.portfolio-item::before {{
  content: "\00b7";
  position: absolute;
  left: 1px;
  color: #888;
  font-weight: bold;
}}
.portfolio-item a {{
  color: #2980b9;
  text-decoration: none;
}}
.portfolio-item a:hover {{
  text-decoration: underline;
}}
/* ---- Self Evaluation ---- */
.self-evaluation {{
  font-size: 12.5px;
  line-height: 1.55;
  color: #444;
}}
.self-eval-item {{
  position: relative;
  padding-left: 12px;
  margin-bottom: 1px;
  font-size: 12.5px;
  line-height: 1.55;
  color: #444;
}}
.self-eval-item::before {{
  content: "\00b7";
  position: absolute;
  left: 1px;
  color: #888;
  font-weight: bold;
}}
.self-eval-label {{
  font-weight: 600;
  color: #222;
}}
/* ---- Print ---- */
@page {{
  size: A4;
  margin: 0;
}}
@media print {{
  body {{ background: white; }}
  .resume {{ margin: 0; padding: 18mm 20mm 14mm; box-shadow: none; max-width: none; }}
  .section {{ page-break-inside: avoid; }}
  .core-advantages {{ page-break-inside: avoid; }}
  .exp-header {{ page-break-after: avoid; }}
  .bullet-list {{ page-break-inside: avoid; }}
  .edu-row {{ page-break-inside: avoid; }}
  .award-row {{ page-break-inside: avoid; }}
  .skill-item {{ page-break-inside: avoid; }}
  .no-print {{ display: none !important; }}
  /* 打印预览提示 */
  .print-hint {{
    display: block;
    background: #fff3cd;
    border: 1px solid #ffc107;
    border-radius: 4px;
    padding: 10px 14px;
    margin-bottom: 12px;
    font-size: 12px;
    color: #856404;
  }}
  .print-hint::before {{
    content: "🖨️ 打印提示：";
    font-weight: bold;
  }}
}}
/* ---- English Layout ---- */
.resume.en .header-left h1 {{
  text-transform: none;
  letter-spacing: 0.5px;
}}
.resume.en .header-contact {{
  font-size: 11.5px;
}}
.resume.en .section-title {{
  font-size: 13px;
  text-transform: uppercase;
  letter-spacing: 0.5px;
}}
.resume.en .edu-school,
.resume.en .exp-company {{
  font-size: 12.5px;
}}
.resume.en .bullet-list li,
.resume.en .skill-item,
.resume.en .core-advantages-list li {{
  font-size: 12px;
}}
</style>
</head>
<body>
<div class="resume{lang_class}" id="resume">
  <div class="print-hint no-print">
    {print_hint_text}
  </div>
  {header_html}
  {core_advantages_html}
  {education_html}
  {experience_html}
  {skills_html}
  {awards_html}
  {portfolio_html}
  {self_evaluation_html}
</div>
<script>
(function() {{
  var el = document.getElementById("resume");
  if (!el) return;
  var maxPages = {max_pages};
  var pageH = 297 - 18 - 14;
  var totalH = pageH * 3.78 * maxPages;
  if (el.scrollHeight > totalH) {{
    el.style.transformOrigin = "top left";
    var scale = totalH / el.scrollHeight;
    el.style.transform = "scale(" + scale.toFixed(4) + ")";
    el.style.width = (100 / scale).toFixed(2) + "%";
  }}
}})();
</script>
</body>
</html>"""


def _esc(text: str) -> str:
    return html_mod.escape(str(text)) if text else ""


def _encode_photo_base64(photo_path: str, max_height: int = 288) -> str:
    if not photo_path or not os.path.isfile(photo_path):
        return ""
    ext = os.path.splitext(photo_path)[1].lower()
    mime_map = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png"}
    mime = mime_map.get(ext, "")
    if not mime:
        return ""
    try:
        with open(photo_path, "rb") as f:
            raw = f.read()
        try:
            from PIL import Image
            import io
            img = Image.open(io.BytesIO(raw))
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            w, h = img.size
            if h > max_height:
                ratio = max_height / h
                img = img.resize((int(w * ratio), max_height), Image.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=92)
            b64 = base64.b64encode(buf.getvalue()).decode("ascii")
            mime = "image/jpeg"
        except ImportError:
            b64 = base64.b64encode(raw).decode("ascii")
        return f"data:{mime};base64,{b64}"
    except Exception:
        return ""


def _build_header_html(data: dict) -> str:
    name = _esc(data.get("name", ""))
    gender = _esc(data.get("gender", ""))
    birth_date = _esc(data.get("birth_date", ""))
    current_city = _esc(data.get("current_city", ""))
    domicile_city = _esc(data.get("domicile_city", ""))
    phone = _esc(data.get("phone", ""))
    email = _esc(data.get("email", ""))
    objective = _esc(data.get("objective", ""))
    target_industry = _esc(data.get("target_industry", ""))
    target_city = _esc(data.get("target_city", ""))
    expected_salary = _esc(data.get("expected_salary", ""))
    job_type = _esc(data.get("job_type", ""))
    photo_path = data.get("photo", "")

    line1_parts = []
    if gender:
        line1_parts.append(f'<span class="label">性别：</span>{gender}')
    if birth_date:
        line1_parts.append(f'<span class="label">出生：</span>{birth_date}')
    if current_city:
        line1_parts.append(f'<span class="label">现居：</span>{current_city}')
    if domicile_city:
        line1_parts.append(f'<span class="label">户籍：</span>{domicile_city}')
    line1 = "&nbsp;&nbsp;|&nbsp;&nbsp;".join(line1_parts)

    line2_parts = []
    if phone:
        line2_parts.append(f'<span class="label">电话：</span>{phone}')
    if email:
        line2_parts.append(f'<span class="label">邮箱：</span>{email}')
    line2 = "&nbsp;&nbsp;|&nbsp;&nbsp;".join(line2_parts)

    line3_parts = []
    if objective:
        line3_parts.append(f'<span class="label">求职意向：</span>{objective}')
    if job_type:
        line3_parts.append(f'<span class="label">类型：</span>{job_type}')
    if target_city:
        line3_parts.append(f'<span class="label">目标城市：</span>{target_city}')
    if target_industry:
        line3_parts.append(f'<span class="label">意向行业：</span>{target_industry}')
    if expected_salary:
        line3_parts.append(f'<span class="label">期望薪资：</span>{expected_salary}')
    line3 = "&nbsp;&nbsp;|&nbsp;&nbsp;".join(line3_parts)

    contact_lines = [l for l in [line1, line2, line3] if l]
    contact_html = "<br>".join(contact_lines)

    photo_b64 = _encode_photo_base64(photo_path) if photo_path else ""
    photo_tag = f'\n  <img class="header-photo" src="{photo_b64}" alt="证件照">' if photo_b64 else ""
    return (
        '<div class="header">\n'
        f'  <div class="header-left"><h1>{name}</h1>\n'
        f'  <div class="header-contact">{contact_html}</div></div>\n'
        f'  {photo_tag}\n'
        "</div>"
    )


def _build_education_html(education) -> str:
    if not education:
        return ""
    if isinstance(education, dict):
        education = [education]
    edu_rows = []
    for edu in education:
        school = _esc(edu.get("school", ""))
        tag = _esc(edu.get("school_tag", ""))
        major = _esc(edu.get("major", ""))
        degree = _esc(edu.get("degree", ""))
        period = _esc(edu.get("period", ""))
        gpa = edu.get("gpa", "")
        courses = edu.get("courses", [])
        tag_html = f'<span class="edu-tag">{tag}</span>' if tag else ""
        detail_parts = [p for p in [major, degree] if p]
        detail = " | ".join(detail_parts)
        gpa_suffix = f" | GPA: {_esc(gpa)}" if gpa else ""
        courses_html = ""
        if courses:
            courses_text = "、".join(_esc(c) for c in courses)
            courses_html = f'<div class="edu-courses">● 主修课程：{courses_text}</div>'
        edu_rows.append(
            f'  <div class="edu-row">\n'
            f'    <div class="edu-left"><span class="edu-school">{school}</span>{tag_html}<span class="edu-detail">&nbsp;|&nbsp;{detail}{gpa_suffix}</span></div>\n'
            f'    <span class="edu-date">{period}</span>\n'
            f'  </div>\n'
            f'  {courses_html}'
        )
    return (
        '<div class="section">\n'
        '  <div class="section-title">教育背景</div>\n'
        + "\n".join(edu_rows) + "\n"
        "</div>"
    )


def _build_experience_html(experience: list) -> str:
    if not experience:
        return ""
    sections = {}
    for item in experience:
        sec = item.get("section", "经历")
        sections.setdefault(sec, []).append(item)
    html_parts = []
    for sec_name, items in sections.items():
        html_parts.append(f'<div class="section">\n  <div class="section-title">{_esc(sec_name)}</div>')
        for item in items:
            company = _esc(item.get("company", ""))
            role = _esc(item.get("role", "") or item.get("title", ""))
            period = _esc(item.get("period", ""))
            bullets = item.get("bullets", [])
            if not bullets:
                for key, label in [("situation", "场景背景"), ("task", "任务职责"), ("action", "行动方法"), ("result", "项目成果")]:
                    val = item.get(key, "")
                    if val:
                        bullets.append({"label": label, "text": val})
            bullet_html = ""
            for b in bullets:
                bl = _esc(b.get("label", ""))
                bt = _esc(b.get("text", ""))
                if bl and bt:
                    bullet_html += f'<li><span class="bullet-label">{bl}：</span>{bt}</li>\n'
                elif bt:
                    bullet_html += f'<li>{bt}</li>\n'
            html_parts.append(
                f'  <div class="exp-header">\n'
                f'    <div class="exp-left"><span class="exp-company">{company}</span><span class="exp-role">{role}</span></div>\n'
                f'    <span class="exp-date">{period}</span>\n'
                f"  </div>\n"
                f'  <ul class="bullet-list">\n{bullet_html}  </ul>'
            )
        html_parts.append("</div>")
    return "\n".join(html_parts)


def _build_skills_html(skills: list) -> str:
    if not skills:
        return ""
    html_parts = ['<div class="section">\n  <div class="section-title">岗位技能</div>']
    for s in skills:
        label = _esc(s.get("label", ""))
        text = _esc(s.get("text", ""))
        if label and text:
            html_parts.append(f'  <div class="skill-item"><span class="skill-label">{label}：</span>{text}</div>')
        elif text:
            html_parts.append(f'  <div class="skill-item">{text}</div>')
    html_parts.append("</div>")
    return "\n".join(html_parts)


def _build_awards_html(awards: list) -> str:
    if not awards:
        return ""
    html_parts = ['<div class="section">\n  <div class="section-title">荣誉奖励</div>']
    for a in awards:
        name = _esc(a.get("name", ""))
        level = _esc(a.get("level", ""))
        time = _esc(a.get("time", ""))
        level_html = f'<span class="award-level">（{level}）</span>' if level else ""
        time_html = f'<span class="award-time">{time}</span>' if time else ""
        html_parts.append(
            f'  <div class="award-row">'
            f'<span><span class="award-name">{name}</span>{level_html}</span>'
            f'{time_html}</div>'
        )
    html_parts.append("</div>")
    return "\n".join(html_parts)


def _build_portfolio_html(portfolio: list) -> str:
    if not portfolio:
        return ""
    html_parts = ['<div class="section">\n  <div class="section-title">作品集</div>']
    for p in portfolio:
        title = _esc(p.get("title", ""))
        url = _esc(p.get("url", ""))
        if title and url:
            html_parts.append(f'  <div class="portfolio-item"><a href="{url}" target="_blank">{title}</a></div>')
        elif title:
            html_parts.append(f'  <div class="portfolio-item">{title}</div>')
    html_parts.append("</div>")
    return "\n".join(html_parts)


def _build_self_evaluation_html(text) -> str:
    if not text:
        return ""
    if isinstance(text, list):
        html_parts = ['<div class="section">\n  <div class="section-title">自我评价</div>']
        for item in text:
            label = _esc(item.get("label", ""))
            desc = _esc(item.get("text", ""))
            if label and desc:
                html_parts.append(f'  <div class="self-eval-item"><span class="self-eval-label">{label}：</span>{desc}</div>')
            elif desc:
                html_parts.append(f'  <div class="self-eval-item">{desc}</div>')
        html_parts.append("</div>")
        return "\n".join(html_parts)
    return (
        '<div class="section">\n'
        '  <div class="section-title">自我评价</div>\n'
        f'  <div class="self-evaluation">{_esc(text)}</div>\n'
        '</div>'
    )


def _build_core_advantages_html(advantages) -> str:
    if not advantages:
        return ""
    if isinstance(advantages, str):
        advantages = [advantages]
    html_parts = ['<div class="core-advantages">\n  <div class="core-advantages-title">核心优势</div>\n  <ul class="core-advantages-list">']
    for adv in advantages:
        html_parts.append(f'    <li>{_esc(adv)}</li>')
    html_parts.append('  </ul>\n</div>')
    return "\n".join(html_parts)


def _detect_language(data: dict) -> str:
    name = data.get("name", "")
    objective = data.get("objective", "")
    import re
    has_cjk = bool(re.search(r'[\u4e00-\u9fff]', name + objective))
    return "" if has_cjk else " en"


def generate_html(resume_data: dict, pages: int = 1) -> str:
    lang_class = _detect_language(resume_data)
    header_html = _build_header_html(resume_data)
    core_advantages_html = _build_core_advantages_html(resume_data.get("core_advantages", []))
    education_html = _build_education_html(resume_data.get("education", {}))
    experience_html = _build_experience_html(resume_data.get("experience", []))
    skills_html = _build_skills_html(resume_data.get("skills", []))
    awards_html = _build_awards_html(resume_data.get("awards", []))
    portfolio_html = _build_portfolio_html(resume_data.get("portfolio", []))
    self_evaluation_html = _build_self_evaluation_html(resume_data.get("self_evaluation", ""))
    if pages == 1:
        print_hint_text = '按 Ctrl+P 打印为 PDF，建议勾选"背景图形"以保留样式。简历内容已优化为 A4 一页。'
    else:
        print_hint_text = f'按 Ctrl+P 打印为 PDF，建议勾选"背景图形"以保留样式。简历内容已优化为 A4 {pages} 页。'
    return HTML_TEMPLATE.format(
        name=_esc(resume_data.get("name", "简历")),
        lang_class=lang_class,
        print_hint_text=print_hint_text,
        max_pages=pages,
        header_html=header_html,
        core_advantages_html=core_advantages_html,
        education_html=education_html,
        experience_html=experience_html,
        skills_html=skills_html,
        awards_html=awards_html,
        portfolio_html=portfolio_html,
        self_evaluation_html=self_evaluation_html,
    )


# ---------------------------------------------------------------------------
# Setup 命令
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# 字段映射与容错
# ---------------------------------------------------------------------------

FIELD_ALIASES = {
    "name": ["姓名", "名字", "fullName", "full_name", "candidate_name"],
    "gender": ["性别", "sex"],
    "birth_date": ["出生日期", "生日", "出生年月", "birthday", "date_of_birth", "dob"],
    "current_city": ["现居城市", "现居地", "居住城市", "current_location", "location"],
    "domicile_city": ["户籍城市", "户籍", "籍贯", "hometown", "domicile"],
    "phone": ["手机", "电话", "联系电话", "联系方式", "mobile", "tel", "telephone", "cell"],
    "email": ["邮箱", "电子邮件", "email_address", "e_mail"],
    "objective": ["求职意向", "意向岗位", "目标岗位", "应聘岗位", "target_position", "position", "job_title"],
    "target_industry": ["意向行业", "目标行业", "industry"],
    "target_city": ["目标城市", "期望城市", "preferred_location"],
    "expected_salary": ["期望薪资", "期望月薪", "salary_expectation", "expected_salary_range"],
    "job_type": ["求职类型", "工作类型", "employment_type", "work_type"],
    "photo": ["照片", "证件照", "头像", "avatar", "profile_photo"],
    "self_evaluation": ["自我评价", "个人评价", "个人总结", "summary", "profile", "about_me"],
    "core_advantages": ["核心优势", "个人优势", "核心竞争力", "key_strengths", "highlights"],
}

EDU_ALIASES = {
    "school": ["学校", "院校", "university", "college", "institution"],
    "degree": ["学历", "学位", "degree_level"],
    "major": ["专业", "专业名称", "field", "subject"],
    "period": ["时间段", "时间", "起止时间", "date", "duration", "graduation_date"],
    "school_tag": ["标签", "学校标签", "school_level", "tag"],
    "gpa": ["绩点", "GPA", "grade_point_average"],
    "courses": ["课程", "主修课程", "core_courses", "relevant_courses"],
}

EXP_ALIASES = {
    "section": ["分类", "类型", "category", "type"],
    "company": ["公司", "单位", "组织", "项目名", "organization", "employer", "project_name"],
    "role": ["职位", "岗位", "角色", "title", "position", "role_title"],
    "period": ["时间段", "时间", "起止时间", "date", "duration"],
}

SKILL_ALIASES = {
    "label": ["类别", "分类", "category", "skill_category"],
    "text": ["描述", "内容", "details", "description"],
}

AWARD_ALIASES = {
    "name": ["奖项", "奖项名称", "award", "title"],
    "level": ["级别", "等级", "rank"],
    "time": ["时间", "获奖时间", "date", "year"],
}


def _map_field(source_key: str, aliases: dict) -> str:
    key_lower = source_key.lower().strip()
    for target, alias_list in aliases.items():
        if key_lower == target.lower():
            return target
        for alias in alias_list:
            if key_lower == alias.lower():
                return target
    return source_key


def _map_dict(data: dict, aliases: dict) -> dict:
    mapped = {}
    used_keys = set()
    for target, alias_list in aliases.items():
        if target in data:
            mapped[target] = data[target]
            used_keys.add(target)
            continue
        for alias in alias_list:
            if alias in data:
                mapped[target] = data[alias]
                used_keys.add(alias)
                break
    for k, v in data.items():
        if k not in used_keys and k not in mapped:
            mapped[k] = v
    return mapped


def map_resume(data: dict) -> dict:
    """将任意结构的简历 JSON 映射为标准模板结构，保留无法映射的字段。"""
    if not isinstance(data, dict):
        return data
    mapped = _map_dict(data, FIELD_ALIASES)

    edu = mapped.get("education", [])
    if isinstance(edu, dict):
        edu = [edu]
    if isinstance(edu, list):
        mapped["education"] = [_map_dict(e, EDU_ALIASES) if isinstance(e, dict) else e for e in edu]

    exp = mapped.get("experience", [])
    if isinstance(exp, dict):
        exp = [exp]
    if isinstance(exp, list):
        mapped["experience"] = [_map_dict(e, EXP_ALIASES) if isinstance(e, dict) else e for e in exp]

    skills = mapped.get("skills", [])
    if isinstance(skills, dict):
        skills = [skills]
    if isinstance(skills, list):
        mapped["skills"] = [_map_dict(s, SKILL_ALIASES) if isinstance(s, dict) else s for s in skills]

    awards = mapped.get("awards", [])
    if isinstance(awards, dict):
        awards = [awards]
    if isinstance(awards, list):
        mapped["awards"] = [_map_dict(a, AWARD_ALIASES) if isinstance(a, dict) else a for a in awards]

    return mapped


def _cmd_map(json_path: str, output_path: str = "") -> int:
    if not os.path.isfile(json_path):
        print(json.dumps({"error": f"JSON 文件不存在: {json_path}"}, ensure_ascii=False))
        return 1
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    mapped = map_resume(data)
    if not output_path:
        base, ext = os.path.splitext(json_path)
        output_path = f"{base}_mapped{ext}"
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(mapped, f, ensure_ascii=False, indent=2)

    diff_report = []
    for key in data:
        mapped_key = _map_field(key, FIELD_ALIASES)
        if mapped_key != key:
            diff_report.append(f"  {key} → {mapped_key}")
    for key in mapped:
        if key not in data:
            diff_report.append(f"  (新增) {key}")

    print(json.dumps({
        "status": "success",
        "message": f"字段映射完成，共处理 {len(data)} 个顶层字段",
        "output_path": os.path.abspath(output_path),
        "field_mapping": diff_report if diff_report else ["  无需映射，字段已匹配"],
    }, ensure_ascii=False, indent=2))
    return 0


def _cmd_setup() -> int:
    import subprocess

    results = {"steps": []}

    deps = [
        ("PyMuPDF", "fitz", "PyMuPDF"),
        ("python-docx", "docx", "python-docx"),
        ("Pillow", "PIL", "Pillow"),
    ]
    missing = []
    for display_name, import_name, pip_name in deps:
        try:
            __import__(import_name)
            results["steps"].append({"step": f"检查 {display_name}", "status": "ok"})
        except ImportError:
            missing.append((display_name, pip_name))
            results["steps"].append({"step": f"检查 {display_name}", "status": "missing"})

    if missing:
        print("正在安装缺失依赖...")
        for display_name, pip_name in missing:
            try:
                subprocess.check_call(
                    [sys.executable, "-m", "pip", "install", pip_name],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
                results["steps"].append({"step": f"安装 {display_name}", "status": "installed"})
            except Exception as e:
                results["steps"].append({"step": f"安装 {display_name}", "status": "failed", "error": str(e)})

    tracker_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "career_tracker.json")
    if not os.path.isfile(tracker_path):
        tracker_template = {
            "version": "1.0",
            "last_updated": "",
            "applications": [],
            "stats": {"total": 0, "by_status": {}, "by_platform": {}, "response_rate": "0%"},
        }
        with open(tracker_path, "w", encoding="utf-8") as f:
            json.dump(tracker_template, f, ensure_ascii=False, indent=2)
        results["steps"].append({"step": "创建 career_tracker.json", "status": "created", "path": tracker_path})
    else:
        results["steps"].append({"step": "career_tracker.json 已存在", "status": "ok"})

    master_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resume_master.json")
    if not os.path.isfile(master_path):
        master_template = dict(BLANK_TEMPLATE)
        master_template["all_objectives"] = []
        master_template["all_self_evaluations"] = {}
        master_template.pop("objective", None)
        master_template.pop("self_evaluation", None)
        with open(master_path, "w", encoding="utf-8") as f:
            json.dump(master_template, f, ensure_ascii=False, indent=2)
        results["steps"].append({"step": "创建 resume_master.json", "status": "created", "path": master_path})
    else:
        results["steps"].append({"step": "resume_master.json 已存在", "status": "ok"})

    results["status"] = "success"
    results["message"] = (
        "初始化完成！\n"
        "下一步：\n"
        "  1. 编辑 resume_master.json 填写你的全部经历\n"
        "  2. 使用 SKILL.md 按 JD 生成定向简历\n"
        "  3. python generate_resume.py build <json> <output.html> 生成 HTML\n"
        "  4. 使用 career-tracker.md 跟踪投递"
    )
    print(json.dumps(results, ensure_ascii=False, indent=2))
    return 0


# ---------------------------------------------------------------------------
# 版本管理
# ---------------------------------------------------------------------------

VERSIONS_DIR = "resume_versions"


def _cmd_version(action: str, json_path: str = "", version_name: str = "") -> int:
    """简历版本管理：保存、列出、对比、回滚。"""
    versions_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), VERSIONS_DIR)

    if action == "save":
        if not json_path or not version_name:
            print(json.dumps({"error": "用法: python generate_resume.py version save <json_file> <version_name>"}, ensure_ascii=False))
            return 1
        if not os.path.isfile(json_path):
            print(json.dumps({"error": f"JSON 文件不存在: {json_path}"}, ensure_ascii=False))
            return 1
        os.makedirs(versions_dir, exist_ok=True)
        from datetime import datetime
        safe_name = version_name.replace("/", "_").replace("\\", "_").replace(":", "_")
        dest = os.path.join(versions_dir, f"{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        data["_version_meta"] = {
            "name": version_name,
            "saved_at": datetime.now().isoformat(),
            "source_file": json_path,
        }
        with open(dest, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(json.dumps({
            "status": "success",
            "message": f"版本已保存: {version_name}",
            "path": dest,
        }, ensure_ascii=False, indent=2))
        return 0

    elif action == "list":
        if not os.path.isdir(versions_dir):
            print(json.dumps({"status": "success", "versions": []}, ensure_ascii=False, indent=2))
            return 0
        versions = []
        for fname in sorted(os.listdir(versions_dir)):
            if fname.endswith(".json"):
                fpath = os.path.join(versions_dir, fname)
                try:
                    with open(fpath, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    meta = data.get("_version_meta", {})
                    versions.append({
                        "file": fname,
                        "name": meta.get("name", fname),
                        "saved_at": meta.get("saved_at", ""),
                        "source": meta.get("source_file", ""),
                    })
                except Exception:
                    versions.append({"file": fname, "name": fname, "saved_at": "", "source": ""})
        print(json.dumps({"status": "success", "versions": versions}, ensure_ascii=False, indent=2))
        return 0

    elif action == "diff":
        if len(sys.argv) < 5:
            print(json.dumps({"error": "用法: python generate_resume.py version diff <version1> <version2>"}, ensure_ascii=False))
            return 1
        v1_name = sys.argv[3]
        v2_name = sys.argv[4]
        if not os.path.isdir(versions_dir):
            print(json.dumps({"error": "没有保存的版本"}, ensure_ascii=False))
            return 1
        v1_path = v2_path = None
        for fname in os.listdir(versions_dir):
            if fname.endswith(".json"):
                try:
                    with open(os.path.join(versions_dir, fname), "r", encoding="utf-8") as f:
                        d = json.load(f)
                    meta = d.get("_version_meta", {})
                    if meta.get("name") == v1_name:
                        v1_path = os.path.join(versions_dir, fname)
                    if meta.get("name") == v2_name:
                        v2_path = os.path.join(versions_dir, fname)
                except Exception:
                    pass
        if not v1_path or not v2_path:
            print(json.dumps({"error": f"找不到版本: {v1_name} 或 {v2_name}，先用 version list 查看可用版本"}, ensure_ascii=False))
            return 1
        with open(v1_path, "r", encoding="utf-8") as f:
            d1 = json.load(f)
        with open(v2_path, "r", encoding="utf-8") as f:
            d2 = json.load(f)
        diff = _diff_resume(d1, d2)
        print(json.dumps({
            "status": "success",
            "version1": v1_name,
            "version2": v2_name,
            "diff": diff,
        }, ensure_ascii=False, indent=2))
        return 0

    elif action == "restore":
        if not version_name:
            print(json.dumps({"error": "用法: python generate_resume.py version restore <version_name> [output_path]"}, ensure_ascii=False))
            return 1
        if not os.path.isdir(versions_dir):
            print(json.dumps({"error": "没有保存的版本"}, ensure_ascii=False))
            return 1
        target_path = None
        for fname in os.listdir(versions_dir):
            if fname.endswith(".json"):
                try:
                    with open(os.path.join(versions_dir, fname), "r", encoding="utf-8") as f:
                        d = json.load(f)
                    meta = d.get("_version_meta", {})
                    if meta.get("name") == version_name:
                        target_path = os.path.join(versions_dir, fname)
                        break
                except Exception:
                    pass
        if not target_path:
            print(json.dumps({"error": f"找不到版本: {version_name}"}, ensure_ascii=False))
            return 1
        output_path = sys.argv[4] if len(sys.argv) >= 5 else version_name.replace("/", "_").replace("\\", "_") + "_restored.json"
        with open(target_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        data.pop("_version_meta", None)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(json.dumps({
            "status": "success",
            "message": f"已回滚到版本: {version_name}",
            "output_path": os.path.abspath(output_path),
        }, ensure_ascii=False, indent=2))
        return 0

    else:
        print(json.dumps({"error": f"未知版本操作: {action}，支持 save / list / diff / restore"}, ensure_ascii=False))
        return 1


def _diff_resume(old: dict, new: dict) -> dict:
    """对比两个简历版本的差异，返回结构化 diff。"""
    diff = {"added": [], "removed": [], "modified": []}

    def _compare_list(old_list, new_list, section_name):
        old_items = {item.get("company", "") + item.get("role", ""): item for item in old_list}
        new_items = {item.get("company", "") + item.get("role", ""): item for item in new_list}
        for key in new_items:
            if key not in old_items:
                diff["added"].append(f"{section_name}: 新增 {key}")
            else:
                old_bullets = [b.get("text", "") for b in old_items[key].get("bullets", [])]
                new_bullets = [b.get("text", "") for b in new_items[key].get("bullets", [])]
                if old_bullets != new_bullets:
                    diff["modified"].append(f"{section_name}: {key} 内容已修改")
        for key in old_items:
            if key not in new_items:
                diff["removed"].append(f"{section_name}: 删除 {key}")

    for section in ["education", "experience", "skills", "awards", "portfolio"]:
        old_sec = old.get(section, [])
        new_sec = new.get(section, [])
        if isinstance(old_sec, dict):
            old_sec = [old_sec]
        if isinstance(new_sec, dict):
            new_sec = [new_sec]
        if old_sec != new_sec:
            _compare_list(old_sec, new_sec, section)

    for field in ["name", "objective", "self_evaluation", "core_advantages"]:
        if old.get(field) != new.get(field):
            if old.get(field) and not new.get(field):
                diff["removed"].append(f"{field}: 已删除")
            elif not old.get(field) and new.get(field):
                diff["added"].append(f"{field}: 新增")
            else:
                diff["modified"].append(f"{field}: 已修改")

    return diff


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> int:
    if len(sys.argv) < 2:
        print(json.dumps({"error": "用法: python generate_resume.py read|build|init|setup|version ..."}, ensure_ascii=False))
        return 1
    command = sys.argv[1]
    if command == "read":
        if len(sys.argv) < 3:
            print(json.dumps({"error": "请提供简历文件路径"}, ensure_ascii=False))
            return 1
        result = read_resume(sys.argv[2])
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return 0 if result.get("status") == "success" else 1
    elif command == "build":
        if len(sys.argv) < 4:
            print(json.dumps({"error": "用法: python generate_resume.py build <json_file> <output_path> [--pages 1|2]"}, ensure_ascii=False))
            return 1
        json_path = sys.argv[2]
        output_path = sys.argv[3]
        pages = 1
        if "--pages" in sys.argv:
            idx = sys.argv.index("--pages")
            if idx + 1 < len(sys.argv):
                try:
                    pages = int(sys.argv[idx + 1])
                    if pages not in (1, 2):
                        print(json.dumps({"error": "--pages 仅支持 1 或 2"}, ensure_ascii=False))
                        return 1
                except ValueError:
                    print(json.dumps({"error": "--pages 参数必须为整数 1 或 2"}, ensure_ascii=False))
                    return 1
        if not os.path.isfile(json_path):
            print(json.dumps({"error": f"JSON 文件不存在: {json_path}"}, ensure_ascii=False))
            return 1
        with open(json_path, "r", encoding="utf-8") as f:
            resume_data = json.load(f)
        resume_data = map_resume(resume_data)
        errors = validate_resume(resume_data)
        if errors:
            print(json.dumps({"error": "简历校验失败", "details": errors, "hint": "字段已自动映射，但仍缺少必填字段。请补充 name/phone/email/objective 后重试"}, ensure_ascii=False, indent=2))
            return 1
        html_content = generate_html(resume_data, pages=pages)
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        print(json.dumps({
            "status": "success",
            "output_path": os.path.abspath(output_path),
            "file_size": len(html_content),
            "pages": pages,
        }, ensure_ascii=False, indent=2))
        return 0
    elif command == "init":
        output_path = sys.argv[2] if len(sys.argv) >= 3 else "resume_template.json"
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(BLANK_TEMPLATE, f, ensure_ascii=False, indent=2)
        print(json.dumps({
            "status": "success",
            "message": "空白简历模板已生成",
            "output_path": os.path.abspath(output_path),
        }, ensure_ascii=False, indent=2))
        return 0
    elif command == "setup":
        return _cmd_setup()
    elif command == "version":
        if len(sys.argv) < 3:
            print(json.dumps({"error": "用法: python generate_resume.py version save|list|diff|restore ..."}, ensure_ascii=False))
            return 1
        action = sys.argv[2]
        json_path = sys.argv[3] if len(sys.argv) >= 4 else ""
        version_name = sys.argv[4] if len(sys.argv) >= 5 else ""
        return _cmd_version(action, json_path, version_name)
    elif command == "map":
        if len(sys.argv) < 3:
            print(json.dumps({"error": "用法: python generate_resume.py map <json_file> [output_path]"}, ensure_ascii=False))
            return 1
        json_path = sys.argv[2]
        output_path = sys.argv[3] if len(sys.argv) >= 4 else ""
        return _cmd_map(json_path, output_path)
    else:
        print(json.dumps({"error": f"未知命令: {command}，支持 read / build / init / setup / version / map"}, ensure_ascii=False))
        return 1


if __name__ == "__main__":
    sys.exit(main())